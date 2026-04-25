# -*- coding: utf-8 -*-
"""
标书资料库 - Word 文档插入模块
支持 python-docx 和 WPS/OFFICE 原生接口
"""

import os, io, tempfile, shutil, uuid
from pathlib import Path

# ── 工具函数 ─────────────────────────────────────────────────────────────────

def _uuid_rid():
    return f"rId{uuid.uuid4().hex[:8]}"

# ── 方案 A: python-docx ──────────────────────────────────────────────────────
# 用于读取 docx、追加段落/图片
#
# 方案 B: WPS/Office COM（仅 Windows）
# 用于完全模拟人工操作（适合复杂格式文档）
#
# 两者并行实现，按需选择

# ── A1: python-docx 追加纯文本段落 ─────────────────────────────────────────

def append_text_paragraph(doc_path, text, style=None, bold=False,
                           font_size=None, color=None, save_path=None):
    """
    用 python-docx 打开已有 docx，追加段落后保存。
    save_path=None 时覆盖原文件。
    """
    import docx
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn

    doc = docx.Document(doc_path)

    para = doc.add_paragraph()
    if style:
        para.style = style

    run = para.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        r, g, b = color
        run.font.color.rgb = RGBColor(r, g, b)

    out = save_path or doc_path
    doc.save(out)
    return out

# ── A2: python-docx 追加图片 ─────────────────────────────────────────────────

def append_image(doc_path, image_path_or_bytes,
                  width_cm=12.0, height_cm=None, save_path=None):
    """
    追加图片到 docx 末尾。
    image_path_or_bytes: 图片文件路径 或 bytes 数据。
    """
    import docx
    from docx.shared import Cm, Inches

    doc = docx.Document(doc_path)

    if isinstance(image_path_or_bytes, (str, os.PathLike)):
        with open(image_path_or_bytes, "rb") as f:
            img_data = f.read()
    else:
        img_data = image_path_or_bytes

    # 计算高度等比缩放
    from PIL import Image as PILImage
    img_buf = io.BytesIO(img_data)
    with PILImage.open(img_buf) as img:
        orig_w, orig_h = img.size
    aspect = orig_h / orig_w
    w_cm = width_cm
    h_cm = height_cm or (width_cm * aspect)

    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(
        io.BytesIO(img_data),
        width=Cm(w_cm),
        height=Cm(h_cm)
    )

    out = save_path or doc_path
    doc.save(out)
    return out

# ── A3: 追加 PDF 附件（嵌入式对象，仅 python-docx）─────────────────────────

def append_pdf_as_object(doc_path, pdf_path, display_name=None, save_path=None):
    """
    在段落中嵌入 PDF（OLE 对象，MS Word / WPS 均支持）。
    """
    import docx
    from docx.oxml.ns import qn
    from lxml import etree
    import zipfile

    display_name = display_name or os.path.basename(pdf_path)

    # python-docx 对 OLE 对象支持有限，改为将 PDF 转为图片组追加
    # （WPS COM 方案见下文）
    return append_pdf_as_images(doc_path, pdf_path, save_path=save_path)

# ── A4: PDF 转图片组追加到 docx ─────────────────────────────────────────────

def append_pdf_as_images(doc_path, pdf_path, dpi=120, page_range=None,
                           width_cm=14.0, save_path=None):
    """
    将 PDF 每页转为图片，依次追加到 docx。
    page_range: None=全部, 或 list如[0,1,2]指定页（从0起）
    """
    import fitz

    pdf_doc = fitz.open(pdf_path)
    out = save_path or doc_path

    tmp_pdf_imgs = []
    for idx in range(pdf_doc.page_count):
        if page_range is not None and idx not in page_range:
            continue
        page = pdf_doc[idx]
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        tmp_pdf_imgs.append(img_bytes)

    pdf_doc.close()

    # 逐页追加（每次重新打开+保存，大文档可用临时文件优化）
    cur = doc_path
    for i, img_bytes in enumerate(tmp_pdf_imgs):
        # 新增分页图注
        from docx.shared import Pt
        import docx
        cur_doc = docx.Document(cur)
        caption_para = cur_doc.add_paragraph()
        caption_run = caption_para.add_run(f"（附图：PDF 第 {i+1} 页）")
        caption_run.font.size = Pt(9)
        caption_run.italic = True
        cur_doc.save(cur)

        cur = append_image(cur, img_bytes, width_cm=width_cm)

    # 最终保存到目标路径
    if save_path and cur != save_path:
        shutil.copy(cur, save_path)
    return out


def _append_pdf_bytes_as_images(doc, pdf_bytes, dpi=120, page_range=None, width_cm=14.0):
    import fitz
    from docx.shared import Cm, Pt

    pdf_buf = io.BytesIO(pdf_bytes)
    with fitz.open(stream=pdf_buf, filetype="pdf") as pdf_doc:
        for idx in range(pdf_doc.page_count):
            if page_range is not None and idx not in page_range:
                continue
            page = pdf_doc[idx]
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")

            caption_para = doc.add_paragraph()
            caption_run = caption_para.add_run(f"（附图：PDF 第 {idx+1} 页）")
            caption_run.font.size = Pt(9)
            caption_run.italic = True
            doc.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))

    return doc

# ── A5: 追加条目文本（支持多段落） ─────────────────────────────────────────

def append_entry_text(doc_path, entry_obj, save_path=None):
    """
    将资料库条目（文本+附件）整体追加到 docx。
    entry_obj: dict with keys raw_text, attachments
    attachments: [{id, file_name, file_ext, file_data, mime_type}]
    """
    import docx
    from docx.shared import Cm, Pt, RGBColor
    import io as _io

    doc = docx.Document(doc_path)

    # 标题（条目名称）
    title_para = doc.add_heading(entry_obj.get("title","未命名条目"), level=2)
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 正文
    raw = entry_obj.get("raw_text", "")
    if raw:
        for para_text in raw.split("\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)

    # 附件
    for att in entry_obj.get("attachments", []):
        file_ext = att.get("file_ext","").lower()
        file_data = att.get("file_data",b"")
        if not file_data:
            continue
        fname = att.get("file_name","附件")

        if file_ext in {".jpg",".jpeg",".png",".gif",".bmp",".webp"}:
            # 图片
            p = doc.add_paragraph()
            r = p.add_run(f"【附件图片：{fname}】")
            r.font.size = Pt(9)
            r.italic = True
            doc.add_picture(_io.BytesIO(file_data), width=Cm(14))
        elif file_ext == ".pdf":
            # PDF 转为图片追加
            p = doc.add_paragraph()
            r = p.add_run(f"【附件PDF：{fname}】")
            r.font.size = Pt(9)
            r.italic = True
            _append_pdf_bytes_as_images(doc, file_data, width_cm=14)
        else:
            # 其他文件仅加文本说明
            p = doc.add_paragraph()
            r = p.add_run(f"【附件：{fname}】")
            r.font.size = Pt(9)
            r.italic = True

    out = save_path or doc_path
    doc.save(out)
    return out

# ── 方案 B: WPS/Office COM（仅 Windows） ────────────────────────────────────
# 适合需要保留文档原有格式的场景（字体、页眉页脚、样式表等）

WIN32_IMPORT = """
try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False
"""

def _has_win32():
    try:
        import win32com.client
        return True
    except ImportError:
        return False

def append_text_com(doc_path, text, bold=False, font_size=12,
                     font_name="宋体", save_path=None):
    """
    通过 WPS/Office COM 接口追加文本（模拟人工输入，保留原文档格式）。
    仅 Windows 可用，自动识别 WPS 或 Office。
    """
    if not _has_win32():
        raise RuntimeError("win32com 未安装，请在 Windows 上安装 pywin32")

    import win32com.client
    import pythoncom

    pythoncom.CoInitialize()

    # 尝试 WPS，失败则用 MS Office
    app = None
    for prog in ["WPS.Application", "KWPS.Application",
                 "Word.Application"]:
        try:
            app = win32com.client.Dispatch(prog)
            break
        except Exception:
            continue

    if app is None:
        pythoncom.CoUninitialize()
        raise RuntimeError("未检测到 WPS 或 Microsoft Office")

    app.Visible = False
    try:
        doc = app.Documents.Open(os.path.abspath(doc_path))
        sel = app.Selection
        sel.EndKey(Unit=6)  # wdStory
        sel.TypeParagraph()

        if bold:
            sel.Font.Bold = True
        sel.Font.Size = font_size
        sel.Font.Name = font_name
        sel.TypeText(text)

        out = save_path or doc_path
        doc.SaveAs(os.path.abspath(out))
        doc.Close()
    finally:
        app.Quit()
        pythoncom.CoUninitialize()

    return out

def insert_at_bookmark_com(doc_path, bookmark_name, text,
                             save_path=None):
    """在书签位置插入文本（WPS/Office COM）"""
    if not _has_win32():
        raise RuntimeError("win32com 未安装")

    import win32com.client, pythoncom

    pythoncom.CoInitialize()
    app = win32com.client.Dispatch("WPS.Application")
    app.Visible = False
    try:
        doc = app.Documents.Open(os.path.abspath(doc_path))
        if bookmark_name in [b.Name for b in doc.Bookmarks]:
            bm = doc.Bookmarks(bookmark_name)
            bm.Select()
            app.Selection.TypeText(text)
        out = save_path or doc_path
        doc.SaveAs(os.path.abspath(out))
        doc.Close()
    finally:
        app.Quit()
        pythoncom.CoUninitialize()
    return out

# ── 批量插入：从资料库选择条目 ───────────────────────────────────────────────

def batch_insert_to_doc(doc_path, entries_data, save_path=None,
                         method="docx"):
    """
    将多个资料库条目批量插入到 Word 文档末尾。
    entries_data: list of entry dicts (含 attachments)
    method: "docx"（推荐）或 "com"（需 Windows + WPS/Office）
    """
    out = doc_path
    for i, entry in enumerate(entries_data):
        if method == "docx":
            out = append_entry_text(out, entry)
        elif method == "com":
            raw = entry.get("raw_text","")
            if raw:
                out = append_text_com(out, raw)
    return out
