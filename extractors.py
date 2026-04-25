# -*- coding: utf-8 -*-
"""
标书资料库 - 从老标书提取内容 v4
支持三种解析模式: line(逐行), page(按页码), heading(按标题)
返回结构化段落 + 图片数据
"""
import os, io, tempfile, shutil, re

# ── DOCX ────────────────────────────────────────────────────────────

def extract_from_docx(path, mode="line", heading_levels=None):
    """提取 docx 文本段落（含样式信息）+ 内嵌图片
    mode: "line"=逐段提取, "heading"=按标题分段
    heading_levels: 指定哪些级别作为章节拆分点，默认 [1]
                   例如 [1,2] 表示 H1 和 H2 都独立成章节
    """
    paragraphs = []
    images = []
    try:
        import docx
        doc = docx.Document(path)
        prev_heading = ""
        all_paras = []
        # 调试：收集所有样式名
        _all_styles = set()
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else "Normal"
            _all_styles.add(style_name)

        # DEBUG：将样式名写入日志文件（兼容 GUI 模式，sys.stdout 可能为 None）
        if mode == "heading":
            try:
                import os as _os
                _log_dir = _os.path.join(_os.environ.get('APPDATA', _os.path.expanduser('~')), 'tender_lib')
                _os.makedirs(_log_dir, exist_ok=True)
                _styles_sorted = sorted(_all_styles)
                _heading_styles = [s for s in _styles_sorted if re.search(r'(标题|Heading)', s, re.IGNORECASE)]
                _msg = f"[DEBUG] 文档样式名: {_styles_sorted}\n[DEBUG] 标题样式: {_heading_styles}\n"
                with open(_os.path.join(_log_dir, 'debug_styles.log'), 'a', encoding='utf-8') as _f:
                    _f.write(_msg + '\n')
            except Exception:
                pass

        # 关键修复：空段落也要记录，因为图片可能在空段落中
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"
            level = 0
            # 检测标题级别 — 支持多种样式名格式：
            #   Word: Heading 1/2/3...
            #   WPS:  标题 1/2/3...  或  toc 1/2/3...  或  目录 1/2/3...
            if text:
                m = re.search(r'(标题|Heading|toc|目录)[\s\-_]*(\d+)', style_name, re.IGNORECASE)
                if m:
                    level = int(m.group(2))
                elif style_name == "Title" or style_name.lower() in ("title", "封面"):
                    level = 1
                elif re.match(r'^(第[一二三四五六七八九十百\d]+条|[\d]+[\.、])', text):
                    level = 4  # 法规/条款样式
            all_paras.append({
                "idx": i, "text": text, "style": style_name,
                "level": level, "is_heading": level > 0,
            })

        if mode == "heading":
            # 按标题分段：每个指定级别的标题及其下的正文合并为一段
            levels = heading_levels if heading_levels else [1]
            paragraphs = _group_by_heading(all_paras, levels)
            # 建立段落索引→条目映射（用于图片关联）
            idx_to_entry = {}
            for entry in paragraphs:
                # 关键修复：每个 entry 必须有 images 字段（_group_by_heading 创建的新 dict 没有此字段）
                entry.setdefault("images", [])
                start_idx = entry.get("start_idx", entry["idx"])
                for i in range(start_idx, entry.get("end_idx", start_idx) + 1):
                    idx_to_entry[i] = entry
            # heading 模式下也要把所有段落原始 idx（尤其是 body）加入映射
            for p in all_paras:
                if p["idx"] not in idx_to_entry:
                    # 找到覆盖此 idx 的条目（通过 body_indices 查找）
                    covering = None
                    for entry in paragraphs:
                        bi = entry.get("body_indices", [])
                        si = entry.get("start_idx", entry["idx"])
                        ei = entry.get("end_idx", si)
                        if si <= p["idx"] <= ei or p["idx"] in bi:
                            covering = entry
                            break
                    idx_to_entry[p["idx"]] = covering
        else:
            paragraphs = all_paras
            idx_to_entry = {}

        # 提取内嵌图片（关联到所属段落/章节）
        # 每个段落关联的图片列表（用 relId 作为 key）
        para_images = {}   # para_idx -> list of image dicts

        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    img_data = rel.target_part.blob
                    ext = os.path.splitext(rel.target_ref)[1].lstrip(".") or "png"
                    fname = f"docx_img_{rel_id.replace('rId','')}.{ext}"

                    # 找到这张图片属于哪个段落（通过 XML 中 rId 的位置）
                    # 策略1：检查段落的 runs 中是否有图片引用
                    img_para_idx = -1
                    for pi, para in enumerate(doc.paragraphs):
                        for run in para.runs:
                            if hasattr(run, "_r") and run._r.xml and rel_id in run._r.xml:
                                img_para_idx = pi
                                break
                        if img_para_idx >= 0:
                            break
                    # 策略2：检查段落整体 XML（有些图片不在 runs 中）
                    if img_para_idx < 0:
                        for pi, para in enumerate(doc.paragraphs):
                            para_xml = para._element.xml if hasattr(para, '_element') else ''
                            if rel_id in para_xml:
                                img_para_idx = pi
                                break
                    # 策略3：查找文档 body XML 中图片出现的段落位置
                    # 使用 lxml etree 精确定位，避免字符串计数误差
                    if img_para_idx < 0:
                        try:
                            from lxml import etree
                            body_el = doc.element.body
                            # 递归查找包含此 relId 的 <w:p> 元素
                            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                            for pi2, p_el in enumerate(body_el.iterchildren('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')):
                                if pi2 >= len(doc.paragraphs):
                                    break
                                p_xml = etree.tostring(p_el, encoding='unicode')
                                if rel_id in p_xml:
                                    img_para_idx = pi2
                                    break
                        except ImportError:
                            # lxml 不可用时回退到字符串计数
                            try:
                                body_xml = doc.element.body.xml
                                img_pos = body_xml.find(rel_id)
                                if img_pos >= 0:
                                    before_pos = body_xml[:img_pos]
                                    para_count = before_pos.count('<w:p ') + before_pos.count('<w:p>')
                                    if para_count < len(doc.paragraphs):
                                        img_para_idx = para_count
                            except Exception:
                                pass
                        except Exception:
                            pass

                    img_dict = {"name": fname, "ext": ext,
                                "data": img_data, "source": "docx",
                                "para_idx": img_para_idx}

                    # 关联到段落
                    if img_para_idx not in para_images:
                        para_images[img_para_idx] = []
                    para_images[img_para_idx].append(img_dict)

                    # 全局列表（保留，给非章节模式用）
                    images.append(img_dict)
                except Exception:
                    pass

        # 图片关联到章节条目（每条条目关联其范围内的图片）
        # 先给每个 entry 初始化空列表
        for entry in paragraphs:
            entry["images"] = []

        # DEBUG：记录图片关联过程
        debug_log = []
        debug_log.append(f"[图片关联] mode={mode}, paragraphs数={len(paragraphs)}, images数={len(images)}")
        if mode == "heading":
            debug_log.append(f"[图片关联] idx_to_entry键数={len(idx_to_entry)}, 范围={min(idx_to_entry.keys())}-{max(idx_to_entry.keys()) if idx_to_entry else 'N/A'}")

        # 使用已构建的 idx_to_entry 映射（heading 模式）或逐个查找（line 模式）
        matched_count = 0
        for img in images:
            pi = img.get("para_idx", -1)
            if pi < 0:
                debug_log.append(f"[图片关联] 图片 {img.get('name')} para_idx={pi} < 0，跳过")
                continue
            target = idx_to_entry.get(pi) if idx_to_entry else None
            if target is None:
                # line 模式：直接用索引找段落
                for entry in paragraphs:
                    if entry.get("idx") == pi:
                        target = entry
                        break
                # 未找到：图片在空段落（line模式被过滤掉）或表格内
                # 找最近的段落条目
                if target is None:
                    # 先按 idx 顺序排序，找距离最近的
                    best_entry = None
                    best_dist = 999999
                    for entry in paragraphs:
                        eidx = entry.get("idx", -1)
                        if eidx < 0:
                            continue  # 跳过表格条目（负idx）
                        dist = abs(eidx - pi)
                        if dist < best_dist:
                            best_dist = dist
                            best_entry = entry
                    if best_entry is not None:
                        target = best_entry
                        debug_log.append(f"[图片关联] 图片 {img.get('name')} para_idx={pi} 重新分配到 idx={best_entry.get('idx')} (距离={best_dist})")
                    else:
                        debug_log.append(f"[图片关联] 图片 {img.get('name')} para_idx={pi} 找不到对应条目")
            if target is not None:
                # 如果目标段落是空段落（无文字），将图片重新分配给最近的非空段落
                # python-docx 的 add_picture() 会创建独立段落，图片所在的段落通常是空的
                target_text = target.get("text", "").strip()
                if not target_text:
                    best_nonempty = None
                    best_dist = 999999
                    target_idx = target.get("idx", pi)
                    for entry in paragraphs:
                        eidx = entry.get("idx", -1)
                        if eidx < 0:
                            continue
                        etxt = entry.get("text", "").strip()
                        if not etxt:
                            continue  # 跳过空段落
                        dist = abs(eidx - target_idx)
                        if dist < best_dist:
                            best_dist = dist
                            best_nonempty = entry
                    if best_nonempty is not None and best_nonempty is not target:
                        debug_log.append(f"[图片关联] 图片 {img.get('name')} 从空段落 idx={target_idx} 重分配到非空段落 idx={best_nonempty.get('idx')} (距离={best_dist})")
                        target = best_nonempty
                target["images"].append(img)
                matched_count += 1

        debug_log.append(f"[图片关联] 成功关联 {matched_count} 张图片")

        # 写入调试日志
        try:
            import os as _os2
            _log_dir2 = _os2.path.join(_os2.environ.get('APPDATA', _os2.expanduser('~')), 'tender_lib')
            _os2.makedirs(_log_dir2, exist_ok=True)
            with open(_os2.path.join(_log_dir2, 'debug_img_assoc.log'), 'a', encoding='utf-8') as _f:
                _f.write('\n'.join(debug_log) + '\n\n')
        except Exception:
            pass
            # 旧格式：para_images 不再使用，直接追加到 entry

        # 提取表格（结构化存储）
        for ti, table in enumerate(doc.tables):
            rows_data = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(cells)
            if rows_data and any(any(row) for row in rows_data):
                paragraphs.append({
                    "idx": -ti - 1,  # 表格用负索引，避免与段落冲突
                    "type": "table",
                    "rows": rows_data,
                    "text": "[表格] " + " | ".join(rows_data[0] if rows_data else []),  # 摘要
                    "style": "Table",
                    "level": 0,
                    "is_heading": False,
                    "images": [],
                })
    except Exception as e:
        paragraphs.append({"idx": 0, "text": f"[DOCX读取失败: {e}]",
                           "style": "Normal", "level": 0, "is_heading": False})
    return paragraphs, images


def _group_by_heading(all_paras, heading_levels=None):
    """
    将段落列表按目录等级分组，每级标题独立成一条条目。
    heading_levels: 要作为章节拆分点的级别列表，默认 [1]
                   例如 [1, 2] 表示 H1 和 H2 都独立成章节
    每条 = 标题 + 其下所有正文（直到下一同级或更高等级标题为止）
    返回每条含 start_idx / end_idx 供图片关联使用。
    """
    if heading_levels is None:
        heading_levels = [1]
    elif isinstance(heading_levels, str):
        heading_levels = [1]
    hl_set = set(heading_levels)

    result = []
    current = None      # 当前正在收集的条目 {"heading", "body_indices", "start_idx"}
    body_buf = []        # 正文行文本列表

    def flush_current():
        """将 current 条目写入 result"""
        if current is None:
            return
        heading_text = current["heading"]["text"].strip()
        body_text = "\n".join(body_buf).strip()
        combined = f"{heading_text}\n{body_text}" if body_text else heading_text
        entry = {
            "idx":       current["heading"]["idx"],
            "start_idx": current["start_idx"],
            "end_idx":   current["body_indices"][-1] if current["body_indices"] else current["heading"]["idx"],
            "text":      combined,
            "style":     current["heading"]["style"],
            "level":     current["heading"]["level"],
            "is_heading": True,
        }
        result.append(entry)

    for p in all_paras:
        lvl  = p.get("level", 0)
        is_h = lvl in hl_set

        if is_h:
            flush_current()
            current = {
                "heading": p,
                "body_indices": [],
                "start_idx": p["idx"],
            }
            body_buf = []
        elif current is not None:
            body = p.get("text", "").strip()
            # 空段落也要入 body_indices，否则图片无法关联（图片常在空段落中）
            current["body_indices"].append(p["idx"])
            if body:
                body_buf.append(body)
        else:
            # 标题之前的零散正文 → 单独一条（不关联章节）
            result.append(p)

    flush_current()
    return result


def _doc_to_docx_via_soffice(path):
    """用 LibreOffice 将 .doc 转成 .docx"""
    soffice = _find_soffice()
    if not soffice:
        return None
    tmp_dir = tempfile.mkdtemp(prefix="tender_doc_")
    try:
        import subprocess
        env = dict(os.environ)
        env.setdefault("LC_ALL", "zh_CN.UTF-8")
        env.setdefault("LANG", "zh_CN.UTF-8")
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx",
             "--outdir", tmp_dir, os.path.abspath(path)],
            env=env, timeout=120, capture_output=True)
        if result.returncode == 0:
            candidates = [f for f in os.listdir(tmp_dir) if f.endswith(".docx")]
            if candidates:
                return os.path.join(tmp_dir, candidates[0])
    except Exception:
        pass
    return None


def extract_from_doc(path, mode="line"):
    """提取 .doc 文本段落 + 图片"""
    paragraphs = []
    images = []
    tmp_docx = _doc_to_docx_via_soffice(path)
    if not tmp_docx:
        paragraphs.append({"idx": 0, "text": "[LibreOffice 未安装，无法处理 .doc 文件]\n"
                                          "[建议将 .doc 另存为 .docx 后导入]",
                           "style": "Normal", "level": 0, "is_heading": False})
        return paragraphs, images
    try:
        paragraphs, images = extract_from_docx(tmp_docx, mode=mode, heading_levels=heading_levels)
    except Exception as e:
        paragraphs.append({"idx": 0, "text": f"[DOC处理失败: {e}]",
                           "style": "Normal", "level": 0, "is_heading": False})
    finally:
        shutil.rmtree(os.path.dirname(tmp_docx), ignore_errors=True)
    return paragraphs, images


# ── PDF ─────────────────────────────────────────────────────────────

def extract_from_pdf(path, max_pages=None, mode="line"):
    """提取 PDF 文本段落 + 每页转图片
    mode: "line"=逐行, "page"=按页码(同页内容合并为一段)
    """
    paragraphs = []
    images = []
    try:
        import fitz
        doc = fitz.open(path)
        total = doc.page_count
        pages_to_scan = min(total, max_pages) if max_pages else total
        for i in range(pages_to_scan):
            page = doc[i]
            text = page.get_text("text").strip()
            page_num = i + 1
            if text:
                if mode == "page":
                    # 按页码分段：整页文本合并为一个段落
                    paragraphs.append({
                        "idx": i, "text": text,
                        "style": "Normal", "level": 0,
                        "is_heading": False, "page": page_num,
                    })
                else:
                    # 逐行分段
                    for line_idx, line in enumerate(text.split("\n")):
                        line = line.strip()
                        if line:
                            paragraphs.append({
                                "idx": i, "text": line,
                                "style": "Normal", "level": 0,
                                "is_heading": False, "page": page_num,
                            })
            # 将该页转成图片
            try:
                mat = fitz.Matrix(1.5, 1.5)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")
                fname = f"pdf_page_{i+1:03d}.png"
                images.append({"name": fname, "ext": "png",
                               "data": img_bytes, "source": "pdf", "page": page_num})
            except Exception:
                pass
        doc.close()
    except Exception as e:
        paragraphs.append({"idx": 0, "text": f"[PDF读取失败: {e}]",
                           "style": "Normal", "level": 0, "is_heading": False})
    return paragraphs, images


# ── TXT ─────────────────────────────────────────────────────────────

def extract_from_txt(path, mode="line"):
    """提取 TXT 文本段落"""
    paragraphs = []
    images = []
    for enc in ("utf-8", "gbk", "gb18030", "utf-16"):
        try:
            with open(path, encoding=enc) as f:
                lines = f.readlines()
            for i, line in enumerate(lines):
                line = line.strip()
                if line:
                    paragraphs.append({"idx": i, "text": line,
                                       "style": "Normal", "level": 0,
                                       "is_heading": False})
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    if not paragraphs:
        paragraphs.append({"idx": 0, "text": "[TXT文件读取失败，编码不支持]",
                           "style": "Normal", "level": 0, "is_heading": False})
    return paragraphs, images


# ── 统一入口 ───────────────────────────────────────────────────────

SUPPORTED_EXT = {".docx", ".doc", ".pdf", ".txt"}

def extract_from_file(path, mode="line", on_progress=None, heading_levels=None):
    """从文件提取内容，返回 (paragraphs, images)
    mode: "line"=逐行(默认), "page"=按页码分段, "heading"=按标题分段
    paragraphs: [{"idx", "text", "style", "level", "is_heading", "page?"}]
    images:     [{"name", "ext", "data", "source", "page?"}]
    on_progress(msg): 可选的进度回调
    """
    ext = os.path.splitext(path)[1].lower()
    basename = os.path.basename(path)

    if on_progress:
        on_progress(f"\u5f00\u59cb\u89e3\u6790: {basename}")

    if ext == ".docx":
        paras, imgs = extract_from_docx(path, mode=mode, heading_levels=heading_levels)
    elif ext == ".doc":
        paras, imgs = extract_from_doc(path, mode=mode)
    elif ext == ".pdf":
        paras, imgs = extract_from_pdf(path, mode=mode)
    elif ext == ".txt":
        paras, imgs = extract_from_txt(path, mode=mode)
    else:
        paras = [{"idx": 0, "text": f"[\u4e0d\u652f\u6301\u7684\u6587\u4ef6\u7c7b\u578b: {ext}]",
                  "style": "Normal", "level": 0, "is_heading": False}]
        imgs = []

    if on_progress:
        on_progress(f"\u63d0\u53d6\u5b8c\u6210: {basename}\uff0c{len(paras)}\u6bb5\u843d\uff0c{len(imgs)}\u5f20\u56fe\u7247")

    return paras, imgs


def _find_soffice():
    """查找 LibreOffice soffice 路径"""
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice 7\program\soffice.exe",
        r"C:\Program Files\LibreOffice 8\program\soffice.exe",
    ]
    for p in candidates:
        if os.path.isfile(p):
            return p
    import shutil
    return shutil.which("soffice")
