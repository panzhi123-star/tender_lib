# -*- coding: utf-8 -*-
"""主窗口 App 类"""
import os, sys, io, traceback, json, re, webbrowser, threading
import tkinter as tk
import tkinter.ttk as ttk
import tkinter.messagebox as mb
import tkinter.filedialog as fd
import tkinter.simpledialog as sd
from tkinter.scrolledtext import ScrolledText
from tkinter import colorchooser
import sqlite3
from PIL import Image, ImageTk

_THIS_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _THIS_DIR not in sys.path:
    sys.path.insert(0, _THIS_DIR)

from config import *
import tender_lib_db2 as db
import extractors
import docx_inserter as inserter
from .entry_dlg import EntryDlg
from .insert_dlg import InsertDlg
from .compare_dlg import CompareDlg
from .import_preview import ImportPreviewDlg
from .keyword_rule import KeywordRuleDlg
from .template import TemplateDlg, TemplateInsertDlg
from .batch_dlg import BatchDlg

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("标书资料库管理工具 v2.7.0")
        self.root.geometry("1280x780")
        self.root.minsize(720, 560)
        self.root.configure(bg=C_BG)

        self._apply_styles()
        db.init_db()
        self.current_cat = None
        self.current_entry = None
        self.img_refs = []
        self.selected_ids = set()
        self.search_var = tk.StringVar(value="")

        self._build_ui()
        self._refresh_cats()
        self._refresh_list()
        self._update_stats()
        self.log_insert("系统就绪。欢迎使用标书资料库 v2.7.0")
        self.root.protocol("WM_DELETE_WINDOW", self.root.destroy)

    # ── 样式 ───────────────────────────────────────────────
    def _apply_styles(self):
        s = ttk.Style()
        try:
            s.theme_use("clam")
        except Exception:
            pass
        s.configure("Treeview", rowheight=28, font=FONT_BODY,
                     background=C_CARD, fieldbackground=C_CARD,
                     foreground=C_TEXT, bordercolor=C_BORDER)
        s.map("Treeview", background=[("selected", "#3498DB")],
              foreground=[("selected", "white")])
        s.configure("Treeview.Heading", font=FONT_HEAD,
                     background=C_BG, foreground=C_TEXT, borderwidth=0)
        s.configure("TNotebook.Tab", font=FONT_BODY, padding=[12, 5])
        s.configure("TNotebook", background=C_PANEL)
        try:
            from ctypes import windll
            windll.shcore.SetProcessDpiAwareness(2)
        except Exception:
            pass

    # ── UI ────────────────────────────────────────────────
    def _build_ui(self):
        # ---- 顶部深色标题栏 ----
        top_bar = tk.Frame(self.root, bg=C_SIDEBAR, height=40)
        top_bar.pack(side=tk.TOP, fill=tk.X)
        top_bar.pack_propagate(False)

        logo_canvas = tk.Canvas(top_bar, width=32, height=32, bg=C_SIDEBAR, highlightthickness=0)
        logo_canvas.pack(side=tk.LEFT, padx=(10, 4), pady=4)
        logo_canvas.create_oval(2, 2, 30, 30, fill="#3498DB", outline="")
        logo_canvas.create_oval(4, 4, 28, 28, fill="#5DADE2", outline="")
        logo_canvas.create_rectangle(9, 10, 25, 24, fill="white", outline="")
        logo_canvas.create_line(9, 14, 25, 14, fill="#3498DB", width=1)
        logo_canvas.create_line(9, 17, 25, 17, fill="#BDC3C7", width=1)
        logo_canvas.create_line(9, 20, 25, 20, fill="#BDC3C7", width=1)

        tk.Label(top_bar, text="标书资料库管理工具  v2.7.0",
                 font=FONT_TITLE, bg=C_SIDEBAR, fg="white", anchor="w").pack(side=tk.LEFT, padx=4, pady=6)
        self.ver_lbl = tk.Label(top_bar, text="", font=FONT_SMALL, bg=C_SIDEBAR, fg=C_SBTXT)
        self.ver_lbl.pack(side=tk.RIGHT, padx=14, pady=6)

        # ---- 工具栏 ----
        toolbar = tk.Frame(self.root, bg=C_CARD2, height=40)
        toolbar.pack(side=tk.TOP, fill=tk.X)
        toolbar.pack_propagate(False)

        for txt, cmd, accent in [
            ("新建", self._new_cat, True),
            ("编辑分类", self._edit_cat, False),
            ("导入标书", self._import_file, True),
            ("插入到Word", self._insert_to_doc, True),
        ]:
            bg_c = C_ACCENT if accent else C_PANEL
            fg_c = "white" if accent else C_TEXT
            b = tk.Button(toolbar, text=txt, command=cmd, font=FONT_BODY, bg=bg_c, fg=fg_c, bd=0, cursor="hand2", padx=12, pady=4, relief=tk.FLAT)
            b.pack(side=tk.LEFT, padx=4, pady=4)

        sep = tk.Frame(toolbar, bg=C_BORDER, width=1)
        sep.pack(side=tk.LEFT, fill=tk.Y, padx=6, pady=4)
        self.stats_lbl = tk.Label(toolbar, text="", font=FONT_SMALL, bg=C_CARD2, fg=C_SUBTXT)
        self.stats_lbl.pack(side=tk.LEFT, padx=8, pady=4)
        for txt, cmd in [("刷新", self._refresh_all), ("导出日志", self._export_logs), ("重置", self._reset_library)]:
            b = tk.Button(toolbar, text=txt, command=cmd, font=FONT_SMALL, bg=C_PANEL, fg=C_TEXT, bd=0, cursor="hand2", padx=8, pady=3, relief=tk.FLAT)
            b.pack(side=tk.RIGHT, padx=4, pady=4)
        
        # 删除选中按钮
        del_btn = tk.Button(toolbar, text="删除选中", 
                           command=self._del_selected,
                           font=FONT_SMALL, bg=C_DANGER, fg="white", 
                           bd=0, cursor="hand2", padx=8, pady=3, relief=tk.FLAT)
        del_btn.pack(side=tk.RIGHT, padx=4, pady=4)

        # ---- 主区域：PanedWindow 三栏 ----
        self.pw = tk.PanedWindow(self.root, sashrelief=tk.FLAT, sashwidth=4, bg=C_BORDER, opaqueresize=True)
        self.pw.pack(fill=tk.BOTH, expand=True)

        # 左栏：分类树
        left = tk.Frame(self.pw, bg=C_SIDEBAR, width=200)
        self.pw.add(left, width=200, stretch="never")
        tree_hdr = tk.Frame(left, bg=C_SIDEBAR, height=32)
        tree_hdr.pack(fill=tk.X)
        tree_hdr.pack_propagate(False)
        tk.Label(tree_hdr, text="分类", font=FONT_HEAD, bg=C_SIDEBAR, fg="white", padx=10).pack(side=tk.LEFT, pady=4)

        # Scrollable category frame (replaces Treeview for richer button-based UI)
        cat_canvas = tk.Canvas(left, bg=C_SIDEBAR, highlightthickness=0)
        cat_vsb = ttk.Scrollbar(left, orient=tk.VERTICAL, command=cat_canvas.yview)
        self.cat_frame = tk.Frame(cat_canvas, bg=C_SIDEBAR)
        self.cat_frame.bind("<Configure>", lambda e: cat_canvas.configure(scrollregion=cat_canvas.bbox("all")))
        cat_canvas.create_window((0, 0), window=self.cat_frame, anchor="nw")
        cat_canvas.configure(yscrollcommand=cat_vsb.set)
        cat_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, pady=4)
        cat_vsb.pack(side=tk.RIGHT, fill=tk.Y)
        # Mouse wheel scrolling
        def _on_cat_wheel(event):
            cat_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        cat_canvas.bind("<MouseWheel>", _on_cat_wheel)
        self.cat_frame.bind("<MouseWheel>", _on_cat_wheel)

        # 中栏：条目列表
        middle = tk.Frame(self.pw, bg=C_BG, width=500)
        self.pw.add(middle, width=500, stretch="never")
        mid_hdr = tk.Frame(middle, bg="#F1F5F9", height=36)
        mid_hdr.pack(fill=tk.X)
        mid_hdr.pack_propagate(False)
        self.mid_lbl = tk.Label(mid_hdr, text="所有资料", font=FONT_HEAD, bg="#F1F5F9", fg=C_TEXT, anchor="w", padx=10)
        self.mid_lbl.pack(side=tk.LEFT, pady=4)
        search_entry = tk.Entry(mid_hdr, textvariable=self.search_var, font=FONT_SMALL, width=16, bd=0, relief=tk.FLAT)
        search_entry.pack(side=tk.RIGHT, padx=4, pady=4)
        search_entry.bind("<Return>", lambda e: self._refresh_list())
        tk.Label(mid_hdr, text="搜索:", font=FONT_SMALL, bg="#F1F5F9", fg=C_SUBTXT).pack(side=tk.RIGHT, padx=(10,0))
        self.sel_count_lbl = tk.Label(mid_hdr, text="", font=FONT_SMALL, bg="#F1F5F9", fg=C_ACCENT)
        self.sel_count_lbl.pack(side=tk.RIGHT, padx=10, pady=4)
        self.entry_tree = ttk.Treeview(middle, style="Entry.Treeview")
        entry_s = ttk.Style()
        entry_s.configure("Entry.Treeview", background=C_PANEL, foreground=C_TEXT, fieldbackground=C_PANEL, rowheight=28, font=FONT_BODY)
        entry_s.map("Entry.Treeview", background=[("selected", "#D6EAF8")], foreground=[("selected", C_ACCENT)])
        self.entry_tree["columns"] = ("名称", "字数", "分类", "标签")
        for col, w in [("名称", 280), ("字数", 70), ("分类", 120), ("标签", 120)]:
            self.entry_tree.heading(col, text=col)
            self.entry_tree.column(col, width=w, anchor="w")
        self.entry_tree.column("#0", width=28, anchor="center")
        entry_vsb = ttk.Scrollbar(middle, orient=tk.VERTICAL, command=self.entry_tree.yview)
        entry_hsb = ttk.Scrollbar(middle, orient=tk.HORIZONTAL, command=self.entry_tree.xview)
        self.entry_tree.configure(yscrollcommand=entry_vsb.set, xscrollcommand=entry_hsb.set)
        self.entry_tree.pack(fill=tk.BOTH, expand=True, padx=4, pady=4)
        entry_vsb.pack(side=tk.RIGHT, fill=tk.Y)
        entry_hsb.pack(side=tk.BOTTOM, fill=tk.X)
        self.entry_tree.bind("<<TreeviewSelect>>", self._on_select)
        self.entry_tree.bind("<Double-Button-1>", self._on_entry_dbl_click)
        self.entry_tree.bind("<Button-3>", self._show_entry_menu)
        self.entry_tree.bind("<Control-a>", lambda e: self._select_all_main())

        # 右栏：内容预览
        right = tk.Frame(self.pw, bg=C_PANEL)
        self.pw.add(right, width=420, stretch="never")
        right_hdr = tk.Frame(right, bg="#F1F5F9", height=36)
        right_hdr.pack(fill=tk.X)
        right_hdr.pack_propagate(False)
        self.right_lbl = tk.Label(right_hdr, text="内容预览", font=FONT_HEAD, bg="#F1F5F9", fg=C_TEXT, anchor="w", padx=10)
        self.right_lbl.pack(side=tk.LEFT, pady=4)
        self.preview_text = ScrolledText(right, wrap=tk.WORD, font=FONT_BODY, bg=C_PANEL, insertbackground=C_ACCENT, padx=10, pady=8, state=tk.DISABLED, highlightthickness=0)
        self.preview_text.pack(fill=tk.BOTH, expand=True, padx=4, pady=4)
        self.img_canvas = tk.Canvas(right, bg=C_CARD2, highlightthickness=0, height=90)
        self.img_hsb = ttk.Scrollbar(right, orient=tk.HORIZONTAL, command=self.img_canvas.xview)
        self.img_inner = tk.Frame(self.img_canvas, bg=C_CARD2)
        self.img_canvas.create_window((0, 0), window=self.img_inner, anchor="nw")
        self.img_inner.bind("<Configure>", lambda e: self.img_canvas.configure(scrollregion=self.img_canvas.bbox("all")))
        self.img_canvas.configure(xscrollcommand=self.img_hsb.set)
        self.img_area_frame = tk.Frame(right, bg=C_CARD2)
        self.img_area_frame.pack(fill=tk.X, padx=4, pady=(0, 4))
        self.img_canvas.pack(in_=self.img_area_frame, fill=tk.X)
        self.img_hsb.pack(in_=self.img_area_frame, side=tk.BOTTOM, fill=tk.X)
        self.img_refs = []

        # ---- 详情笔记本（附件/版本/模板）----
        detail_nb = ttk.Notebook(right)
        detail_nb.pack(fill=tk.BOTH, expand=True, padx=4, pady=(0, 4))

        # 附件标签页
        att_frame = tk.Frame(detail_nb, bg=C_PANEL)
        detail_nb.add(att_frame, text=" 附件 ")
        self.atree = ttk.Treeview(att_frame, columns=("文件名","类型","大小"),
                                   show="headings", height=4)
        for col, w in [("文件名",200),("类型",60),("大小",80)]:
            self.atree.heading(col, text=col)
            self.atree.column(col, width=w, anchor="w")
        self.atree.pack(fill=tk.BOTH, expand=True)

        # 版本标签页
        ver_frame = tk.Frame(detail_nb, bg=C_PANEL)
        detail_nb.add(ver_frame, text=" 版本 ")
        self.vtree = ttk.Treeview(ver_frame, columns=("版本","时间","字数"),
                                   show="headings", height=4)
        for col, w in [("版本",60),("时间",140),("字数",60)]:
            self.vtree.heading(col, text=col)
            self.vtree.column(col, width=w, anchor="w")
        self.vtree.pack(fill=tk.BOTH, expand=True)

        # 模板标签页
        tpl_frame = tk.Frame(detail_nb, bg=C_PANEL)
        detail_nb.add(tpl_frame, text=" 模板 ")
        self.tpl_tree = ttk.Treeview(tpl_frame, columns=("名称","分类"),
                                      show="headings", height=4)
        for col, w in [("名称",200),("分类",120)]:
            self.tpl_tree.heading(col, text=col)
            self.tpl_tree.column(col, width=w, anchor="w")
        self.tpl_tree.pack(fill=tk.BOTH, expand=True)

        # ---- 日志面板 ----
        log_frame = tk.Frame(self.root, bg="#1E272E", height=100)
        log_frame.pack(side=tk.BOTTOM, fill=tk.X)
        log_frame.pack_propagate(False)
        log_hdr = tk.Frame(log_frame, bg="#2C3E50", height=22)
        log_hdr.pack(fill=tk.X)
        log_hdr.pack_propagate(False)
        tk.Label(log_hdr, text="操作日志", font=FONT_SMALL, bg="#2C3E50", fg="#AEB6BF",
                 anchor="w", padx=8).pack(side=tk.LEFT)
        self.log_text = tk.Text(log_frame, wrap=tk.WORD, font=("Consolas", 9),
                                bg="#1E272E", fg="#7DFFAF", insertbackground="#7DFFAF",
                                height=5, bd=0, padx=6, pady=4)
        log_vsb = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=log_vsb.set)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        log_vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.config(state="disabled")

        # ---- 状态栏 ----
        status = tk.Frame(self.root, bg=C_PANEL, height=26)
        status.pack(side=tk.BOTTOM, fill=tk.X)
        status.pack_propagate(False)
        self.status_lbl = tk.Label(status, text="就绪", font=FONT_SMALL, bg=C_PANEL, fg=C_SUBTXT, anchor="w", padx=10)
        self.status_lbl.pack(side=tk.LEFT, pady=2)

        self.pw.sash_place(0, 198, 0)
        self.pw.sash_place(1, 696, 0)
        self.search_var = tk.StringVar()
        self.selected_ids = set()
        self._refresh_all()

    def _refresh_list(self):
        """更新左侧列表"""
        for row in self.entry_tree.get_children():
            self.entry_tree.delete(row)
        cat_id = self.current_cat
        kw = self.search_var.get().strip()
        if cat_id == -1:
            # Uncategorised: show entries with category_id IS NULL
            conn = db.get_conn()
            rows = conn.execute(
                "SELECT * FROM entries WHERE category_id IS NULL "
                "ORDER BY updated_at DESC LIMIT 200").fetchall()
            conn.close()
            rows = [dict(r) for r in rows]
        elif kw:
            rows = db.search_entries(kw, cat_id=cat_id)
        else:
            rows = db.list_entries(cat_id=cat_id)
        # Build category name lookup
        _cat_map = {c["id"]: c["name"] for c in db.list_all_categories()}
        for e in rows:
            char_count = len(e.get("raw_text") or "")
            cat_name = _cat_map.get(e.get("category_id"), "") or ""
            tag_str = e.get("tags") or ""
            self.entry_tree.insert("", "end", iid=str(e["id"]), values=(
                e["title"], char_count, cat_name, tag_str))

    def _refresh_cats(self):
        """Refresh category list with subcategory support and action buttons."""
        for w in self.cat_frame.winfo_children():
            w.destroy()
        
        # 获取所有分类
        all_cats = db.list_all_categories()
        
        # 构建父子关系映射
        parent_map = {}  # parent_id -> [children]
        for c in all_cats:
            pid = c["parent_id"]
            if pid not in parent_map:
                parent_map[pid] = []
            parent_map[pid].append(c)
        
        # 未分类和全部条目按钮
        btn_row = tk.Frame(self.cat_frame, bg=C_SIDEBAR)
        btn_row.pack(fill=tk.X, padx=3, pady=2)
        
        uncategorized_btn = tk.Button(btn_row, text="未分类",
                                     font=FONT_SMALL, bg=C_SBAR2, fg=C_SBTXT,
                                     bd=0, padx=8, pady=3, cursor="hand2",
                                     activebackground=C_ACCENT, activeforeground="white",
                                     command=lambda: self._select_cat(-1))
        uncategorized_btn.pack(side=tk.LEFT, padx=2)
        
        all_btn = tk.Button(btn_row, text="全部条目",
                           font=FONT_SMALL, bg=C_SBAR2, fg=C_SBTXT,
                           bd=0, padx=8, pady=3, cursor="hand2",
                           activebackground=C_ACCENT, activeforeground="white",
                           command=lambda: self._select_cat(None))
        all_btn.pack(side=tk.LEFT, padx=2)
        
        # 递归显示分类
        def show_cats(parent_id, level=0):
            cats = [c for c in all_cats if c["parent_id"] == parent_id]
            for c in cats:
                sel = c["id"] == self.current_cat
                indent = "  " * level
                
                # 创建分类行容器
                row = tk.Frame(self.cat_frame, bg=C_PANEL if not sel else C_ACCENT)
                row.pack(fill=tk.X, padx=3, pady=1)
                
                # 分类名称按钮
                btn = tk.Button(row, text=f"{indent}{c['name']}",
                               font=FONT_BODY, 
                               bg=C_ACCENT if sel else C_PANEL,
                               fg="white" if sel else C_TEXT,
                               bd=0, padx=8, pady=3, anchor="w",
                               cursor="hand2",
                               command=lambda cid=c["id"]: self._select_cat(cid))
                btn.pack(side=tk.LEFT, fill=tk.X, expand=True)
                
                # 操作按钮容器（始终显示）
                btns = tk.Frame(row, bg=C_PANEL if not sel else C_ACCENT)
                btns.pack(side=tk.RIGHT, padx=2)
                
                # 添加子分类按钮（仅一级分类显示）
                if level == 0:
                    add_sub = tk.Button(btns, text="+", font=("Arial", 9, "bold"),
                                       bg="#4CAF50", fg="white", bd=0,
                                       padx=4, pady=0, width=2,
                                       cursor="hand2",
                                       command=lambda cid=c["id"]: self._new_subcat(cid))
                    add_sub.pack(side=tk.LEFT, padx=1)
                    # Tooltip
                    self._create_tooltip(add_sub, "添加子分类")
                
                # 编辑按钮
                edit_btn = tk.Button(btns, text="E", font=("Arial", 9, "bold"),
                                    bg="#2196F3", fg="white", bd=0,
                                    padx=4, pady=0, width=2,
                                    cursor="hand2",
                                    command=lambda cid=c["id"], name=c["name"]: 
                                        self._rename_cat(cid, name))
                edit_btn.pack(side=tk.LEFT, padx=1)
                self._create_tooltip(edit_btn, "重命名")
                
                # 删除按钮
                del_btn = tk.Button(btns, text="X", font=("Arial", 9, "bold"),
                                   bg=C_DANGER, fg="white", bd=0,
                                   padx=4, pady=0, width=2,
                                   cursor="hand2",
                                   command=lambda cid=c["id"]: self._del_cat_direct(cid))
                del_btn.pack(side=tk.LEFT, padx=1)
                self._create_tooltip(del_btn, "删除")
                
                # 右键菜单和双击重命名
                btn.bind("<Button-3>", lambda e, cid=c["id"], name=c["name"]:
                        self._show_cat_menu(e, cid, name))
                btn.bind("<Double-Button-1>", lambda e, cid=c["id"], name=c["name"]:
                        self._rename_cat(cid, name))
                
                # 递归显示子分类
                show_cats(c["id"], level + 1)
        
        # 顶部「添加一级分类」按钮
        top_btn = tk.Button(self.cat_frame, text="+ 添加一级分类",
                           font=FONT_BODY, bg=C_ACCENT, fg="white",
                           bd=0, padx=8, pady=4, cursor="hand2",
                           activebackground=C_ACCENT2, activeforeground="white",
                           command=self._new_cat)
        top_btn.pack(fill=tk.X, padx=6, pady=(6, 4))

        # 从根分类开始显示
        show_cats(None, 0)
    
    def _create_tooltip(self, widget, text):
        """Create a simple tooltip for a widget."""
        def show_tip(e):
            tip = tk.Toplevel(widget)
            tip.wm_overrideredirect(True)
            tip.wm_geometry(f"+{e.x_root+10}+{e.y_root+10}")
            label = tk.Label(tip, text=text, bg="#333", fg="white",
                           font=("Arial", 9), padx=4, pady=2)
            label.pack()
            widget._tooltip = tip
        def hide_tip(e):
            if hasattr(widget, "_tooltip"):
                widget._tooltip.destroy()
                del widget._tooltip
        widget.bind("<Enter>", show_tip)
        widget.bind("<Leave>", hide_tip)
    
    def _new_subcat(self, parent_id):
        """Create a new subcategory under the given parent."""
        name = sd.askstring("新建子分类", "请输入子分类名称：")
        if name and name.strip():
            cid = db.create_category(name.strip(), parent_id=parent_id)
            if cid:
                self._refresh_cats()
                self.status_lbl.config(text=f"已创建子分类：【{name.strip()}】")
            else:
                mb.showinfo("提示", "分类名称已存在")
    
    def _del_cat_direct(self, cid):
        """Delete a category directly by ID."""
        cats = {c["id"]: c for c in db.list_all_categories()}
        if cid not in cats:
            return
        cat = cats[cid]
        # 检查是否有子分类
        children = [c for c in cats.values() if c["parent_id"] == cid]
        if children:
            mb.showinfo("提示", "该分类下有子分类，请先删除子分类")
            return
        # 检查是否有条目
        entries = db.list_entries(cat_id=cid)
        if entries:
            msg = f"分类【{cat['name']}】下有 {len(entries)} 个条目，\n删除分类将把这些条目移至'未分类'，继续吗？"
            if not mb.askyesno("确认", msg):
                return
            # 将条目移至未分类
            for e in entries:
                db.update_entry(e["id"], category_id=None)
        msg2 = f"确定要删除分类【{cat['name']}】吗？"
        if mb.askyesno("确认删除", msg2):
            db.delete_category(cid)
            if self.current_cat == cid:
                self.current_cat = None
            self._refresh_cats()
            self._refresh_list()
            self.status_lbl.config(text=f"已删除分类：【{cat['name']}】")


    def _reset_library(self):
        """一键清空资料库所有入库内容"""
        import sqlite3 as _sq
        from tkinter import messagebox as _mb
        if not _mb.askyesno("确认重置",
            "即将清空资料库中所有入库内容！\n\n"
            "包括：所有条目、分类、附件、关键词规则、版本历史\n\n"
            "此操作不可撤销，确定继续吗？",
            icon="warning"):
            return
        if not _mb.askyesno("二次确认",
            "这是最后确认！所有数据将被永久删除。\n确定吗？"):
            return
        try:
            import tender_lib_db2 as _db2
            db_path = _db2.DB_PATH
            conn = _sq.connect(db_path)
            cur = conn.cursor()
            # 按依赖顺序清空所有表
            for tbl in ["history", "entry_versions", "entry_attachments", "keyword_rules",
                         "entries", "templates", "categories"]:
                try:
                    cur.execute(f"DELETE FROM {tbl}")
                except Exception:
                    pass
            # 重置自增序列
            for tbl in ["history", "entry_versions", "entry_attachments", "keyword_rules",
                         "entries", "templates", "categories"]:
                try:
                    cur.execute(f"DELETE FROM sqlite_sequence WHERE name='{tbl}'")
                except Exception:
                    pass
            conn.commit()
            conn.close()
            self._refresh_all()
            _mb.showinfo("重置完成", "资料库已清空。")
        except Exception as e:
            _mb.showerror("重置失败", str(e))

    def _refresh_all(self):
        """手动刷新首页：重新加载分类和列表"""
        self._refresh_cats()
        self._refresh_list()
        self._update_stats()
        self.log_insert("[刷新] 首页已刷新")
        self.status_lbl.config(text="首页已刷新")

    def _select_cat(self, cat_id):
        self.current_cat = None if cat_id in (None, -1) else cat_id
        self._refresh_cats()
        self._refresh_list()

    def _on_cat_select(self, event=None):
        """Handle category selection (now button-based, kept for compatibility)."""
        pass

    def _clear_search(self):
        self.search_var.set("")
        self._refresh_list()

    def _show_progress(self, text="", maximum=100, mode="determinate"):
        self.status_lbl.config(text=text)
        self.root.update_idletasks()

    def _update_progress(self, value, text=None):
        if text:
            self.status_lbl.config(text=text)
        self.root.update_idletasks()

    def _hide_progress(self, text="就绪"):
        self.status_lbl.config(text=text)

    def _on_select(self, event=None):
        """Track selected entries (supports multi-select)."""
        sel = self.entry_tree.selection()
        self.selected_ids = set(int(s) for s in sel)
        if sel:
            self.current_entry = db.get_entry(int(sel[0]))
            self._show_detail(self.current_entry)
            self._refresh_versions(self.current_entry["id"])
            self.sel_count_lbl.config(
                text=f"已选中 {len(self.selected_ids)} 项")
        else:
            self.current_entry = None
            self.sel_count_lbl.config(text="")

    def _on_entry_dbl_click(self, event):
        """Handle double-click on entry - open edit dialog."""
        self._edit_entry()

    def _show_detail(self, entry):
        self.preview_text.config(state="normal")
        self.preview_text.delete("1.0", tk.END)
        self.preview_text.insert("1.0", entry.get("raw_text", "")[:3000] if entry else "(无内容)")
        self.preview_text.config(state="disabled")
        if entry:
            self._refresh_imgs(entry["id"])
            self._refresh_atts(entry["id"])

    def _refresh_imgs(self, eid):
        if not hasattr(self, 'img_area_frame'):
            return
        for w in self.img_inner.winfo_children():
            w.destroy()
        self.img_refs = []
        atts = [a for a in db.list_attachments(eid)
                if a["file_ext"].lower() in {".png",".jpg",".jpeg",".gif",".bmp",".webp"}]
        if not atts:
            self.img_area_frame.pack_forget()
            return
        self.img_area_frame.pack(fill=tk.X, padx=4, pady=(0, 4))  # 有图片，显示区域
        for att in atts:
            data = db.get_attachment_data(att["id"])
            if not data:
                continue
            try:
                img = Image.open(io.BytesIO(data))
                img.thumbnail((160, 120))
                photo = ImageTk.PhotoImage(img)
                self.img_refs.append(photo)
                fr = tk.Frame(self.img_inner, bg=C_CARD2, bd=1, relief=tk.SOLID)
                tk.Label(fr, image=photo, bg=C_CARD2).pack(padx=4, pady=4)
                tk.Label(fr, text=att["file_name"], font=FONT_SMALL,
                         bg=C_CARD2, fg=C_SUBTXT).pack(pady=0)
                # Move buttons
                btn_fr = tk.Frame(fr, bg=C_CARD2)
                btn_fr.pack(pady=2)
                up_btn = tk.Button(btn_fr, text="↑", font=("Arial", 8), bg=C_ACCENT, fg="white",
                                  bd=0, cursor="hand2", width=2,
                                  command=lambda a=att, ei=eid: self._move_img(ei, a["id"], -1))
                up_btn.pack(side=tk.LEFT, padx=1)
                dn_btn = tk.Button(btn_fr, text="↓", font=("Arial", 8), bg=C_ACCENT, fg="white",
                                  bd=0, cursor="hand2", width=2,
                                  command=lambda a=att, ei=eid: self._move_img(ei, a["id"], 1))
                dn_btn.pack(side=tk.LEFT, padx=1)
                del_btn = tk.Button(btn_fr, text="×", font=("Arial", 8), bg=C_DANGER, fg="white",
                                   bd=0, cursor="hand2", width=2,
                                   command=lambda a=att, ei=eid: self._del_img(ei, a["id"]))
                del_btn.pack(side=tk.LEFT, padx=1)
                fr.pack(side=tk.LEFT, padx=4, pady=8)
            except Exception:
                pass

    def _move_img(self, eid, att_id, direction):
        atts = db.list_attachments(eid)
        img_atts = [a for a in atts if a["file_ext"].lower() in {".png",".jpg",".jpeg",".gif",".bmp",".webp"}]
        ids = [a["id"] for a in img_atts]
        if att_id not in ids:
            return
        idx = ids.index(att_id)
        new_idx = idx + direction
        if new_idx < 0 or new_idx >= len(ids):
            return
        db.swap_attachment_order(ids[idx], ids[new_idx])
        self._refresh_imgs(eid)
        self.log_insert("[图片] 已调整顺序")

    def _del_img(self, eid, att_id):
        if mb.askyesno("确认", "确定删除此图片附件？"):
            db.delete_attachment(att_id)
            self._refresh_imgs(eid)
            self._update_stats()
            self.log_insert("[图片] 已删除附件")

    def _refresh_atts(self, eid):
        for row in self.atree.get_children():
            self.atree.delete(row)
        atts = [a for a in db.list_attachments(eid)
                if a["file_ext"].lower() not in {".png",".jpg",".jpeg",".gif",".bmp",".webp"}]
        for a in atts:
            data = db.get_attachment_data(a["id"])
            sz = len(data) if data else 0
            self.atree.insert("", tk.END, values=(
                a["file_name"],
                a["file_ext"].upper(),
                f"{sz/1024:.1f}KB" if sz < 1048576 else f"{sz/1048576:.1f}MB",
            ), tags=(str(a["id"]),))

    def _refresh_versions(self, eid):
        for row in self.vtree.get_children():
            self.vtree.delete(row)
        vers = db.get_versions(eid)
        for v in vers:
            words = len(v.get("raw_text", "") or "") if v.get("raw_text") else 0
            self.vtree.insert("", tk.END, values=(
                f"v{v['version_no']}",
                v["created_at"][:16],
                str(words),
            ), tags=(str(v["id"]),))

    def _fill_tpl_tree(self):
        for row in self.tpl_tree.get_children():
            self.tpl_tree.delete(row)
        for t in db.list_templates():
            self.tpl_tree.insert("", tk.END, values=(t["name"], t["category"]))

    def _update_stats(self):
        s = db.get_stats()
        self.stats_lbl.config(
            text=f"  条目：{s['total_entries']}  |  分类：{s['total_cats']}  |  "
                 f"附件：{s['total_attachs']}  |  模板：{s['total_templates']}  |  版本：{s['total_versions']}"
        )

    def log_insert(self, msg):
        """追加工作日志（现代风格）"""
        import datetime
        ts = datetime.datetime.now().strftime("%H:%M:%S")
        line = f"[{ts}] {msg}\n"
        self.log_text.config(state="normal")
        # 根据消息类型着色
        if "[错误]" in msg or "[警告]" in msg:
            self.log_text.insert(tk.END, f"[{ts}] ", "ts")
            self.log_text.insert(tk.END, msg + "\n", "err")
        elif "[完成]" in msg or "[成功]" in msg:
            self.log_text.insert(tk.END, f"[{ts}] ", "ts")
            self.log_text.insert(tk.END, msg + "\n", "ok")
        elif "->" in msg:
            self.log_text.insert(tk.END, f"[{ts}] ", "ts")
            self.log_text.insert(tk.END, msg + "\n", "detail")
        else:
            self.log_text.insert(tk.END, f"[{ts}] {msg}\n", "info")
        self.log_text.tag_config("ts", foreground="#5DADE2")
        self.log_text.tag_config("err", foreground="#FF6B6B")
        self.log_text.tag_config("ok", foreground="#7DFFAF")
        self.log_text.tag_config("detail", foreground="#A8FF78")
        self.log_text.tag_config("info", foreground="#7DFFAF")
        self.log_text.see(tk.END)
        self.log_text.config(state="disabled")


    def _export_logs(self):
        """Export debug logs to file."""
        import datetime
        log_content = self.log_text.get("1.0", tk.END)
        if not log_content.strip():
            mb.showinfo("提示", "日志为空，无需导出")
            return
        fp = fd.asksaveasfilename(
            defaultextension=".log",
            filetypes=[("Log files", "*.log"), ("Text files", "*.txt"), ("All files", "*.*")],
            initialfile=f"tender_lib_debug_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        )
        if fp:
            try:
                with open(fp, 'w', encoding='utf-8') as f:
                    f.write(log_content)
                mb.showinfo("成功", f"日志已导出到\n{fp}")
            except Exception as e:
                mb.showerror("错误", f"导出失败: {e}")

    # ── 操作 ───────────────────────────────────────────────
    def _new_entry(self):
        dlg = EntryDlg(self.root, "新建条目", {})
        if dlg.result:
            db.create_entry(**dlg.result)
            self._refresh_list()
            self._update_stats()
            self.status_lbl.config(text=f"已创建：「{dlg.result['title']}」")

    def _edit_entry(self):
        if not self.current_entry:
            mb.showinfo("提示", "请先在列表中选择条目")
            return
        dlg = EntryDlg(self.root, "编辑条目", self.current_entry)
        if dlg.result:
            db.update_entry(self.current_entry["id"], **dlg.result)
            self._refresh_list()
            e = db.get_entry(self.current_entry["id"])
            self.current_entry = e
            self._show_detail(e)
            self._refresh_versions(e["id"])
            self._update_stats()

    def _del_entry(self):
        if not self.current_entry:
            return
        if mb.askyesno("确认", f"确定删除「{self.current_entry['title']}」？"):
            db.delete_entry(self.current_entry["id"])
            self.current_entry = None
            self._refresh_list()
            self._update_stats()
            self.preview_text.config(state="normal")
            self.preview_text.delete("1.0", tk.END)
            self.preview_text.config(state="disabled")

    def _select_all_main(self):
        """全选：选中列表中所有可见条目"""
        kids = list(self.entry_tree.get_children())
        if not kids:
            return
        for kid in kids:
            self.entry_tree.selection_add(kid)
        self.selected_ids = set(int(kid) for kid in kids)
        self.sel_count_lbl.config(
            text=f"\u5df2\u9009\u4e2d {len(self.selected_ids)} \u9879",
            fg=C_ACCENT)
        self.log_insert(f"[\u5168\u9009] \u5df2\u9009\u4e2d {len(kids)} \u9879")

    def _deselect_all_main(self):
        """取消全选"""
        self.entry_tree.selection_remove(*self.entry_tree.selection())
        self.selected_ids = set()
        self.sel_count_lbl.config(text="")
        self.log_insert("[\u5168\u9009] \u5df2\u53d6\u6d88\u9009\u62e9")

    def _del_selected(self):
        """Delete all selected entries (multi-select)."""
        if not self.selected_ids:
            mb.showinfo("提示", "请在列表中用Ctrl/Shift多选要删除的条目")
            return
        n = len(self.selected_ids)
        if not mb.askyesno("确认删除",
                f"确定删除选中的 {n} 个条目？此操作不可撤销。"):
            return
        for eid in list(self.selected_ids):
            db.delete_entry(eid)
        self.selected_ids = set()
        self.current_entry = None
        self._refresh_list()
        self._update_stats()
        self.preview_text.config(state="normal")
        self.preview_text.delete("1.0", tk.END)
        self.preview_text.config(state="disabled")
        self.sel_count_lbl.config(text="")
        self.log_insert(f"[批量删除] 已删除 {n} 个条目")
        self.status_lbl.config(text=f"已删除 {n} 个条目")


    def _new_cat(self):
        name = sd.askstring("新建分类", "输入分类名称：")
        if name and name.strip():
            cid = db.create_category(name.strip())
            if cid:
                self._refresh_cats()
                self._update_stats()
                self.status_lbl.config(text=f"已创建分类：「{name.strip()}」")
            else:
                mb.showinfo("提示", "分类已存在")

    def _edit_cat(self):
        """Edit selected category (rename)."""
        if not self.current_cat:
            mb.showinfo("提示", "请先在左侧点击选择要编辑的分类")
            return
        cats = {c["id"]: c["name"] for c in db.list_categories()}
        old_name = cats.get(self.current_cat, "?")
        new_name = sd.askstring("编辑分类", "输入新名称：", initialvalue=old_name)
        if new_name and new_name.strip() and new_name.strip() != old_name:
            if db.rename_category(self.current_cat, new_name.strip()):
                self._refresh_cats()
                self.status_lbl.config(text=f"已修改分类：「{new_name.strip()}」")
            else:
                mb.showinfo("提示", "该名称已存在或无效")


    # ---- Debug Logging ----
    def _export_debug_log(self, msg):
        import datetime, os
        log_dir = os.path.join(os.environ.get("APPDATA", "."), "tender_lib")
        os.makedirs(log_dir, exist_ok=True)
        log_file = os.path.join(log_dir, "debug.log")
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        with open(log_file, "a", encoding="utf-8") as f:
            f.write(f"[{ts}] {msg}\r\n")

    def _open_debug_log(self):
        import os, subprocess
        log_file = os.path.join(os.environ.get("APPDATA", "."), "tender_lib", "debug.log")
        if os.path.exists(log_file):
            os.startfile(log_file)
        else:
            from tkinter import messagebox
            messagebox.showinfo("日志", "暂无日志文件，先触发一些操作后再导出。")
    def _del_cat(self):
        if not self.current_cat:
            mb.showinfo("提示", "请先在左侧点击选择要删除的分类")
            return
        cats = {c["id"]: c["name"] for c in db.list_categories()}
        name = cats.get(self.current_cat, "?")
        if mb.askyesno("确认", f"删除分类「{name}」？\n条目标记为未分类，不会删除。"):
            db.delete_category(self.current_cat)
            self.current_cat = None
            self._refresh_cats()
            self._refresh_list()
            self._update_stats()

    def _show_cat_menu(self, event, cat_id, cat_name):
        """Show right-click context menu for category."""
        menu = tk.Menu(self.root, tearoff=0, bg=C_PANEL, fg=C_TEXT,
                      activebackground=C_ACCENT, activeforeground="white")
        menu.add_command(label="重命名", 
                        command=lambda: self._rename_cat(cat_id, cat_name))
        menu.add_command(label="删除", 
                        command=lambda: self._delete_cat_by_id(cat_id, cat_name))
        menu.post(event.x_root, event.y_root)

    def _show_entry_menu(self, event):
        """Show right-click context menu for entry tree."""
        menu = tk.Menu(self.root, tearoff=0, bg=C_PANEL, fg=C_TEXT,
                      activebackground=C_ACCENT, activeforeground="white")
        menu.add_command(label="编辑条目", command=self._edit_entry)
        menu.add_command(label="删除条目", command=self._del_entry)
        menu.add_separator()
        menu.add_command(label="插入到Word", command=self._insert_sel)
        menu.post(event.x_root, event.y_root)

    def _rename_cat(self, cat_id, old_name):
        """Rename a category."""
        new_name = sd.askstring("重命名分类", "输入新名称：", initialvalue=old_name)
        if new_name and new_name.strip() and new_name.strip() != old_name:
            if db.rename_category(cat_id, new_name.strip()):
                self._refresh_cats()
                self.status_lbl.config(text=f"已重命名为：「{new_name.strip()}」")
            else:
                mb.showinfo("提示", "该名称已存在")

    def _delete_cat_by_id(self, cat_id, cat_name):
        """Delete category by ID (for context menu)."""
        if mb.askyesno("确认删除", f"删除分类「{cat_name}」？\n\n该分类下的条目将被标记为未分类，不会被删除。"):
            db.delete_category(cat_id)
            if self.current_cat == cat_id:
                self.current_cat = None
                self._refresh_list()
            self._refresh_cats()
            self._update_stats()
            self.status_lbl.config(text=f"已删除分类：「{cat_name}」")

    def _import_file(self):
        """Import with threaded parsing + progress bar."""
        files = fd.askopenfilenames(title="导入标书文件",
            filetypes=[("标书文件", "*.docx *.doc *.pdf *.txt"), ("全部", "*.*")])
        if not files:
            return

        replace_mode = None
        paths_to_preview = []
        for path in files:
            fname = os.path.basename(path)
            existing_eid = db.find_entry_by_source(str(path))
            if existing_eid and replace_mode is None:
                existing_e = db.get_entry(existing_eid)
                replace = mb.askyesnocancel(
                    "文件已存在",
                    f"「{fname}」已在资料库中。\n已有条目：「{existing_e['title']}」\n"
                    f"更新日期：{existing_e['updated_at'][:16]}\n\n"
                    "[是]替换  [否]跳过  [取消]停止",
                    parent=self.root)
                if replace is None:
                    self.log_insert("[取消] 用户取消导入")
                    return
                replace_mode = replace
            if existing_eid and not replace_mode:
                self.log_insert(f"[跳过] {fname} 已存在")
                continue
            if existing_eid and replace_mode:
                db.delete_entry(existing_eid)
                self.log_insert(f"[替换] {fname} 旧条目已删除")
            paths_to_preview.append(path)

        if not paths_to_preview:
            mb.showinfo("提示", "所有文件已在资料库中，全部跳过。")
            return

        # ─── 解析模式选择对话框 ───
        parse_mode = ["line"]
        has_pdf = any(os.path.splitext(p)[1].lower() == ".pdf" for p in paths_to_preview)
        has_docx = any(os.path.splitext(p)[1].lower() in (".docx", ".doc") for p in paths_to_preview)
        if has_pdf or has_docx:
            _md = tk.Toplevel(self.root)
            _md.title("选择解析方式")
            _md.geometry("460x280")
            _md.transient(self.root)
            _md.grab_set()
            _md.resizable(False, False)
            _md.configure(bg=C_PANEL)
            tk.Label(_md, text="请选择标书解析方式：", font=FONT_HEAD,
                     bg=C_PANEL, fg=C_TEXT).pack(anchor="w", padx=24, pady=(24, 12))
            _mv = tk.StringVar(value="line")
            for val, txt, desc in [
                ("line", "逐行解析", "每行文本作为一个独立段落（默认）"),
                ("page", "按页码分段", "同一页的内容合并为一个段落（适用PDF）"),
                ("heading", "按章节分段", "按目录等级拆分为独立章节（需选级别）"),
            ]:
                _fr = tk.Frame(_md, bg=C_PANEL)
                _fr.pack(fill=tk.X, padx=36, pady=5)
                tk.Radiobutton(_fr, text=txt, variable=_mv, value=val,
                               font=FONT_BODY, bg=C_PANEL,
                               activebackground=C_PANEL).pack(side=tk.LEFT)
                tk.Label(_fr, text=f"  — {desc}", font=FONT_SMALL,
                         bg=C_PANEL, fg=C_SUBTXT).pack(side=tk.LEFT)
            # 子对话框：选择章节级别
            def _show_level_dialog():
                _sd = tk.Toplevel()
                _sd.title("选择目录等级")
                _sd.geometry("280x240")
                _sd.transient(self.root)
                _sd.grab_set()
                _sd.configure(bg=C_PANEL)
                tk.Label(_sd, text="请选择作为章节的目录等级：",
                         font=FONT_BODY, bg=C_PANEL, fg=C_TEXT).pack(pady=(16, 10))
                _level_var = tk.IntVar(value=1)
                for lvl, label in [(1, "一级标题 (Heading 1)"), (2, "二级标题 (Heading 2)"),
                                   (3, "三级标题 (Heading 3)"), (4, "其他标题")]:
                    tk.Radiobutton(_sd, text=label, variable=_level_var, value=lvl,
                                   font=FONT_BODY, bg=C_PANEL,
                                   activebackground=C_PANEL).pack(anchor="w", padx=30, pady=2)
                tk.Label(_sd, text="单选，选择拆分章节的目录等级",
                         font=FONT_SMALL, bg=C_PANEL, fg=C_SUBTXT).pack(pady=(4, 8))

                def _sd_ok():
                    _heading_levels[0][:] = [_level_var.get()]
                    parse_mode[0] = "heading"
                    _confirmed[0] = True
                    _sd.destroy()

                tk.Button(_sd, text="确认", command=_sd_ok,
                          font=FONT_BODY, bg=C_ACCENT, fg="white", bd=0,
                          cursor="hand2", padx=18, pady=4).pack(pady=8)
                self.root.wait_window(_sd)

            _confirmed = [False]
            _heading_levels = [[1]]

            def _ok():
                mode = _mv.get()
                if mode == "heading":
                    _pending[0] = True
                    _md.destroy()
                else:
                    parse_mode[0] = mode
                    _confirmed[0] = True
                    _md.destroy()
            def _cc(): _md.destroy()

            _pending = [False]
            _bf = tk.Frame(_md, bg=C_PANEL); _bf.pack(pady=18)
            tk.Button(_bf, text="确认", command=_ok, font=FONT_BODY,
                     bg=C_ACCENT, fg="white", bd=0, cursor="hand2",
                     padx=20, pady=4).pack(side=tk.LEFT, padx=8)
            tk.Button(_bf, text="取消", command=_cc, font=FONT_BODY,
                     bd=0, cursor="hand2", padx=20, pady=4).pack(side=tk.LEFT, padx=8)
            self.root.wait_window(_md)
            if _pending[0]:
                # 弹出级别选择子对话框
                _pending[0] = False
                _show_level_dialog()
            if not _confirmed[0]:
                self.log_insert("[取消] 用户取消导入")
                return

        n = len(paths_to_preview)
        self.log_insert(f"[开始] 正在解析 {n} 个文件...")

        self.status_lbl.config(text=f"正在解析 {n} 个文件...")
        self.root.update_idletasks()

        results_lock = threading.Lock()
        all_paras = []
        all_imgs = []
        errors = []
        prog_lock = threading.Lock()
        prog_done = [0]
        prog_max = n

        def _on_progress(delta):
            with prog_lock:
                prog_done[0] += delta
                p = int(prog_done[0] / prog_max * 100)
            self._update_progress(p, text=f"解析中 {prog_done[0]}/{prog_max}...")
            self.root.update_idletasks()

        def _on_result(path, paras, imgs, fname):
            with results_lock:
                if paras:
                    all_paras.extend(paras)
                    all_imgs.extend(imgs)
                    self.log_insert(f"[解析] {fname} -> {len(paras)}段落,{len(imgs)}图片")
                else:
                    self.log_insert(f"[空文件] {fname} 无可提取内容")

        def _worker():
            mode = parse_mode[0]
            levels = _heading_levels[0] if mode == "heading" else None
            for path in paths_to_preview:
                fname = os.path.basename(path)
                try:
                    paras, imgs = extractors.extract_from_file(path, mode=mode, heading_levels=levels)
                    _on_result(path, paras, imgs, fname)
                except Exception as ex:
                    errors.append((fname, str(ex)))

        t = threading.Thread(target=_worker, daemon=True)
        t.start()

        prog_win = tk.Toplevel(self.root)
        prog_win.title("解析进度")
        prog_win.geometry("460x140")
        prog_win.transient(self.root)
        prog_win.grab_set()
        tk.Label(prog_win, text=f"正在解析 {n} 个文件，请稍候...",
                 font=FONT_BODY, bg=C_PANEL).pack(pady=(16, 8))
        bar_var = tk.DoubleVar(value=0)
        bar = ttk.Progressbar(prog_win, variable=bar_var, maximum=100, length=400)
        bar.pack(padx=20)
        lbl = tk.Label(prog_win, text=f"0/{n}", font=FONT_SMALL,
                       fg=C_SUBTXT, bg=C_PANEL)
        lbl.pack(pady=4)

        def _poll():
            if t.is_alive():
                with prog_lock:
                    done = prog_done[0]
                bar_var.set(int(done / prog_max * 100))
                lbl.config(text=f"{done}/{n}")
                self.root.update_idletasks()
                prog_win.after(120, _poll)
            else:
                prog_win.destroy()
                self._hide_progress()
                if errors:
                    err_msg = "\n".join(f"{f}: {e}" for f, e in errors[:5])
                    mb.showwarning("解析警告",
                        f"部分文件解析出错：\n{err_msg}", parent=self.root)
                if not all_paras:
                    mb.showwarning("提示", "所有文件解析失败，无内容可导入。",
                                   parent=self.root)
                    return
                self.log_insert(f"[预览] {len(all_paras)}段落已解析，打开预览...")
                self.root.update_idletasks()
                self.root.state('normal')
                rules = db.list_keyword_rules()
                try:
                    print("DEBUG: paras=" + str(len(all_paras)) + " imgs=" + str(len(all_imgs)) + " rules=" + str(len(rules)))
                    dlg = ImportPreviewDlg(
                        self.root, all_paras, all_imgs,
                        ", ".join(os.path.basename(p) for p in paths_to_preview),
                        rules)
                    print("DEBUG: ImportPreviewDlg created OK")
                    self.root.wait_window(dlg)
                except Exception as e:
                    err = traceback.format_exc()
                    print("DEBUG: ImportPreviewDlg FAILED: " + str(e))
                    _log_dir = os.path.join(os.environ.get('APPDATA', '.'), 'tender_lib')
                    _err_log = os.path.join(_log_dir, '_dialog_error.log')
                    try:
                        os.makedirs(_log_dir, exist_ok=True)
                        with open(_err_log, 'a', encoding='utf-8') as lf:
                            lf.write('[' + str(os.times()) + '] ImportPreviewDlg failed: ' + str(e) + chr(10) + err + chr(10))
                    except:
                        pass
                    mb.showerror('Error', 'Preview window failed: ' + str(e))
                    self.status_lbl.config(text="Preview failed")
                    return
                self._refresh_list()
                self._update_stats()
                self.log_insert("[完成] 导入预览窗口已关闭")
                self.status_lbl.config(text="导入完成")

        prog_win.after(200, _poll)


    def _batch_import(self):
        """Batch import with threaded processing + progress bar."""
        parse_mode = ["line"]
        folder = fd.askdirectory(title="选择要导入的文件夹")
        if not folder:
            return
        files = []
        for root_d, _, fnames in os.walk(folder):
            for f in fnames:
                ext = os.path.splitext(f)[1].lower()
                if ext in {".docx",".doc",".pdf",".txt"}:
                    files.append(os.path.join(root_d, f))
        if not files:
            mb.showinfo("提示", "该文件夹中没有找到标书文件")
            return

        n = len(files)
        self.log_insert(f"[批量] 开始导入文件夹，共 {n} 个文件...")
        self.status_lbl.config(text=f"正在批量导入 {n} 个文件...")
        self.root.update_idletasks()

        prog_win = tk.Toplevel(self.root)
        prog_win.title("批量导入进度")
        prog_win.geometry("460x140")
        prog_win.transient(self.root)
        prog_win.grab_set()
        tk.Label(prog_win, text=f"正在批量导入 {n} 个文件...",
                 font=FONT_BODY, bg=C_PANEL).pack(pady=(16, 8))
        bar_var = tk.DoubleVar(value=0)
        bar = ttk.Progressbar(prog_win, variable=bar_var, maximum=n, length=400)
        bar.pack(padx=20)
        lbl = tk.Label(prog_win, text=f"0/{n}", font=FONT_SMALL,
                       fg=C_SUBTXT, bg=C_PANEL)
        lbl.pack(pady=4)

        lock = threading.Lock()
        done_count = [0]
        cnt = [0]
        img_total = [0]
        skip = [0]
        errors = []

        def _on_done(fname, status, n_paras, n_imgs):
            with lock:
                dc = done_count[0]
            if status == "skip":
                self.log_insert(f"[{dc}/{n}] [跳过] {fname} 已存在")
            elif status == "ok":
                self.log_insert(f"[{dc}/{n}] {fname} -> {n_paras}段落,{n_imgs}图片")
            elif status == "empty":
                self.log_insert(f"[{dc}/{n}] [空] {fname}")
            elif status == "error":
                self.log_insert(f"[{dc}/{n}] [错误] {fname}")

        def _on_progress():
            with lock: done_count[0] += 1
            bar_var.set(done_count[0])
            lbl.config(text=f"{done_count[0]}/{n}")
            prog_win.update_idletasks()

        def _worker():
            for path in files:
                fname = os.path.basename(path)
                try:
                    existing_eid = db.find_entry_by_source(str(path))
                    if existing_eid:
                        with lock: skip[0] += 1
                        _on_done(fname, "skip", 0, 0)
                        _on_progress()
                        continue
                    paras, imgs = extractors.extract_from_file(path, mode=parse_mode[0])
                    text = "\n".join(p["text"] for p in paras)
                    if paras:
                        eid = db.create_entry(
                            title=fname, raw_text=text,
                            content_type="composite" if imgs else "text",
                            source_file=str(path))
                        for img in imgs:
                            db.add_attachment(eid, img["name"], "."+img["ext"], img["data"])
                        with lock:
                            cnt[0] += 1
                            img_total[0] += len(imgs)
                        _on_done(fname, "ok", len(paras), len(imgs))
                    else:
                        _on_done(fname, "empty", 0, 0)
                except Exception as ex:
                    errors.append(ex)
                    _on_done(fname, "error", 0, 0)
                finally:
                    _on_progress()

        t = threading.Thread(target=_worker, daemon=True)
        t.start()

        def _poll():
            if t.is_alive():
                prog_win.after(150, _poll)
            else:
                prog_win.destroy()
                self._hide_progress()
                self._refresh_list()
                self._update_stats()
                self.log_insert(
                    f"[完成] 共 {cnt[0]} 条目，{img_total[0]} 图片，{skip[0]} 跳过")
                self.status_lbl.config(text=f"批量导入：{cnt[0]}条目")
                if errors:
                    mb.showwarning("部分错误",
                        f"批量导入完成，但有 {len(errors)} 个文件出错。\n"
                        "详情请查看工作日志。")
                else:
                    mb.showinfo("批量导入完成",
                        f"共处理 {cnt[0]} 个文件，{img_total[0]} 张图片。")

        prog_win.after(200, _poll)


    def _add_image(self):
        if not self.current_entry:
            mb.showinfo("提示", "请先在列表中选择条目")
            return
        files = fd.askopenfilenames(title="选择图片",
            filetypes=[("图片", "*.png *.jpg *.jpeg *.bmp *.gif *.webp"),
                       ("全部", "*.*")])
        if not files:
            return
        eid = self.current_entry["id"]
        for path in files:
            with open(path, "rb") as f:
                data = f.read()
            ext = os.path.splitext(path)[1].lower()
            db.add_attachment(eid, os.path.basename(path), ext, data, "image/"+ext.lstrip("."))
        db.update_entry(eid, content_type="image")
        e = db.get_entry(eid)
        self.current_entry = e
        self._show_detail(e)
        self._update_stats()
        mb.showinfo("完成", f"已添加 {len(files)} 张图片")

    def _add_pdf(self):
        if not self.current_entry:
            mb.showinfo("提示", "请先在列表中选择条目")
            return
        path = fd.askopenfilename(title="选择PDF",
            filetypes=[("PDF文件", "*.pdf"), ("全部", "*.*")])
        if not path:
            return
        with open(path, "rb") as f:
            data = f.read()
        db.add_attachment(self.current_entry["id"],
            os.path.basename(path), ".pdf", data, "application/pdf")
        db.update_entry(self.current_entry["id"], content_type="pdf")
        e = db.get_entry(self.current_entry["id"])
        self.current_entry = e
        self._show_detail(e)
        self._update_stats()
        mb.showinfo("完成", "PDF 已添加为附件")

    def _save_att(self):
        sel = self.atree.selection()
        if not sel:
            return
        aid = int(self.atree.item(sel[0])["tags"][0])
        fname = self.atree.item(sel[0])["values"][0]
        data = db.get_attachment_data(aid)
        if not data:
            return
        out = fd.asksaveasfilename(title="保存附件",
            defaultextension=os.path.splitext(fname)[1], initialfile=fname)
        if out:
            with open(out, "wb") as f:
                f.write(data)
            mb.showinfo("保存成功", f"已保存到：{out}")

    def _compare_version(self):
        sel = self.vtree.selection()
        if not sel:
            return
        vid = int(sel[0])
        if not self.current_entry:
            return
        CompareDlg(self.root, db.get_entry(self.current_entry["id"]), vid)

    def _insert_to_doc(self):
        """Insert selected entries into Word with threading + progress."""
        target = fd.askopenfilename(title="选择Word文档（支持WPS）",
            filetypes=[("Word文档", "*.docx"), ("Word 97-2003", "*.doc"), ("全部", "*.*")])
        if not target:
            return
        dlg = InsertDlg(self.root, db.list_entries(limit=300))
        self.root.wait_window(dlg)
        if not dlg.selected_ids:
            return

        entries = []
        for eid in dlg.selected_ids:
            e = db.get_entry(eid)
            e["attachments"] = [dict(a, file_data=db.get_attachment_data(a["id"]))
                                 for a in db.list_attachments(eid)]
            entries.append(e)

        n = len(entries)
        self.log_insert(f"[插入] 准备插入 {n} 个条目到Word...")
        self.status_lbl.config(text=f"正在插入 {n} 个条目...")
        self.root.update_idletasks()

        prog_win = tk.Toplevel(self.root)
        prog_win.title("插入进度")
        prog_win.geometry("420x140")
        prog_win.transient(self.root)
        prog_win.grab_set()
        tk.Label(prog_win, text=f"正在插入 {n} 个条目到Word文档，请稍候...",
                 font=FONT_BODY, bg=C_PANEL).pack(pady=(16, 8))
        bar = ttk.Progressbar(prog_win, mode="indeterminate", length=380)
        bar.pack(padx=20)
        lbl = tk.Label(prog_win, text="正在处理...", font=FONT_SMALL,
                       fg=C_SUBTXT, bg=C_PANEL)
        lbl.pack(pady=4)
        bar.start(12)

        result = [None, None]
        poll_count = [0]  # timeout counter

        def _worker():
            try:
                out = inserter.batch_insert_to_doc(target, entries)
                result[0] = out
            except Exception as ex:
                result[1] = str(ex)

        threading.Thread(target=_worker, daemon=True).start()

        def _poll():
            poll_count[0] += 1
            if result[0] is None and result[1] is None:
                if poll_count[0] > 600:  # 120s timeout
                    result[1] = "操作超时（120秒），请检查Word文档是否被锁定"
                else:
                    prog_win.after(200, _poll)
                    return
            else:
                bar.stop()
                prog_win.destroy()
                self._hide_progress()
                if result[1]:
                    mb.showerror("插入失败", result[1])
                    self.log_insert(f"[错误] 插入失败: {result[1]}")
                else:
                    out = result[0]
                    db.increment_use(dlg.selected_ids)
                    self._refresh_list()
                    self.log_insert(f"[完成] 已插入 {n} 个条目到：{out}")
                    self.status_lbl.config(text="插入完成")
                    mb.showinfo("插入完成",
                        f"已插入 {n} 个条目到：\n{out}")
                    try:
                        os.startfile(out)
                    except Exception:
                        pass

        prog_win.after(200, _poll)

    def _insert_sel(self):
        """Insert currently selected single entry into Word."""
        if not self.current_entry:
            mb.showinfo("提示", "请先在列表中选择条目")
            return
        target = fd.askopenfilename(title="选择Word文档（支持WPS）",
            filetypes=[("Word文档", "*.docx"), ("全部", "*.*")])
        if not target:
            return

        e = self.current_entry
        e["attachments"] = [dict(a, file_data=db.get_attachment_data(a["id"]))
                             for a in db.list_attachments(e["id"])]

        self.log_insert(f"[插入] 正在插入「{e['title']}」...")
        self.status_lbl.config(text=f"正在插入「{e['title']}」...")
        self.root.update_idletasks()

        prog_win = tk.Toplevel(self.root)
        prog_win.title("插入进度")
        prog_win.geometry("400x120")
        prog_win.transient(self.root)
        prog_win.grab_set()
        tk.Label(prog_win, text=f"正在插入「{e['title']}」到Word...",
                 font=FONT_BODY, bg=C_PANEL).pack(pady=(16, 8))
        bar = ttk.Progressbar(prog_win, mode="indeterminate", length=360)
        bar.pack(padx=20)
        bar.start(15)

        def _on_done_ok(out):
            prog_win.destroy()
            self._hide_progress()
            self._refresh_list()
            self.log_insert(f"[完成] 已插入「{e['title']}」")
            self.status_lbl.config(text="插入完成")
            mb.showinfo("插入完成", f"已插入「{e['title']}」到：\n{out}")
            try:
                os.startfile(out)
            except Exception:
                pass

        def _on_done_err(msg):
            prog_win.destroy()
            self._hide_progress()
            mb.showerror("插入失败", msg)
            self.log_insert(f"[错误] 插入失败: {msg}")

        def _worker():
            try:
                out = inserter.batch_insert_to_doc(target, [e])
                db.increment_use([e["id"]])
                self.root.after(0, lambda: _on_done_ok(out))
            except Exception as err:
                self.root.after(0, lambda: _on_done_err(str(err)))

        threading.Thread(target=_worker, daemon=True).start()


    def _open_batch(self):
        BatchDlg(self.root, self)
    def _open_kw_rules(self):
        KeywordRuleDlg(self.root, self)


    def _open_templates(self):
        TemplateDlg(self.root, self)

    def _new_template(self):
        TemplateDlg(self.root, self, mode="new")

    def _insert_from_template(self):
        target = fd.askopenfilename(title="选择Word文档（支持WPS）",
            filetypes=[("Word文档", "*.docx"), ("全部", "*.*")])
        if not target:
            return
        dlg = TemplateInsertDlg(self.root)
        self.root.wait_window(dlg)
        if not dlg.result:
            return
        try:
            out = inserter.batch_insert_to_doc(target, dlg.result)
            mb.showinfo("插入完成", f"模板章节已插入到：\n{out}")
            try:
                os.startfile(out)
            except Exception:
                pass
        except Exception as e:
            mb.showerror("插入失败", str(e))

    # ── 共享模式 ───────────────────────────────────────────
    def _set_share_path(self):
        path = fd.askdirectory(title="选择共享资料库文件夹（局域网共享目录）")
        if path:
            db.set_setting("share_path", path)
            db.set_setting("share_mode", "1")
            mb.showinfo("设置完成",
                f"共享路径已设置为：\n{path}\n\n其他用户打开「打开共享资料库」即可访问。")

    def _open_shared(self):
        path = db.get_setting("share_path")
        if not path:
            mb.showinfo("提示", "请先设置共享路径（菜单：共享 > 设置共享路径）")
            return
        if not os.path.isdir(path):
            mb.showerror("路径无效", f"共享路径不存在：\n{path}")
            return
        # 切换数据库到共享路径
        new_db = os.path.join(path, "tender_library.db")
        if not os.path.exists(new_db):
            mb.showinfo("提示", "共享目录中未找到数据库，将创建新库。")
        # 更新环境变量，通知重启
        mb.showinfo("共享模式",
            f"数据库路径：\n{new_db}\n\n请重启程序以加载共享资料库。")

    def _show_share_status(self):
        mode = db.get_setting("share_mode")
        path = db.get_setting("share_path")
        if mode == "1" and path:
            mb.showinfo("共享状态",
                f"模式：共享资料库\n路径：{path}")
        else:
            mb.showinfo("共享状态", "模式：本地资料库（默认）")

    # ── 备份 / 导出 ───────────────────────────────────────
    def _export_db(self):
        out = fd.asksaveasfilename(title="导出资料库为Word",
            defaultextension=".docx", filetypes=[("Word文档", "*.docx")])
        if not out:
            return
        try:
            import docx
            doc = docx.Document()
            doc.add_heading("标书资料库导出", level=1)
            for e in db.list_entries(limit=500):
                cat = db.list_categories()
                cat_map = {c["id"]: c["name"] for c in cat}
                doc.add_heading(e["title"], level=2)
                doc.add_paragraph(f"分类：{cat_map.get(e.get('category_id'), '')}")
                doc.add_paragraph(e.get("raw_text", "") or "(无内容)")
                doc.add_page_break()
            doc.save(out)
            mb.showinfo("导出成功", f"已保存到：\n{out}")
            try:
                os.startfile(out)
            except Exception:
                pass
        except Exception as e:
            mb.showerror("导出失败", str(e))

    def _backup_db(self):
        out = fd.asksaveasfilename(title="备份数据库",
            defaultextension=".db", filetypes=[("SQLite数据库", "*.db"), ("全部", "*.*")])
        if not out:
            return
        try:
            import shutil
            shutil.copy2(db.DB_PATH, out)
            mb.showinfo("备份成功", f"数据库已备份到：\n{out}")
        except Exception as e:
            mb.showerror("备份失败", str(e))

    def _show_help(self):
        txt = """
【标书资料库管理工具 v2.2 使用说明】

一、资料库管理
  - 新建条目：录入公司资质、业绩等资料
  - 导入文件：批量导入老标书（.docx/.doc/.pdf/.txt）
  - 编辑/删除：对现有条目进行修改或删除

二、图片和PDF附件
  - 在右侧「图片附件」Tab添加资质证书扫描件
  - 在右侧「文件附件」Tab添加PDF文档

三、版本历史
  - 每次编辑都会自动保存历史版本
  - 在「版本历史」Tab查看所有版本并对比差异
  - 双击版本可对比当前内容与历史版本的差异

四、模板管理
  - 创建标书模板，关联常用条目
  - 从模板快速插入多个条目到Word

五、插入Word/WPS
  - 选择目标Word文档
  - 勾选要插入的条目
  - 一键插入，保留原文档格式

六、多人共享
  - 将资料库放在局域网共享文件夹
  - 其他用户通过「共享 > 打开共享资料库」访问
"""
        win = tk.Toplevel(self)
        win.title("使用说明")
        win.geometry("560x500")
        t = ScrolledText(win, wrap=tk.WORD, font=FONT_BODY)
        t.insert("1.0", txt.strip())
        t.config(state="disabled")
        t.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

    def _show_about(self):
        mb.showinfo("关于",
            "标书资料库管理工具 v2.1\n\n"
            "支持：条目管理 / 版本对比 / 模板管理\n"
            "     图片PDF附件 / 批量导入 / 插入Word(WPS)\n"
            "     多人共享资料库\n\n"
            "技术栈：Python + Tkinter + SQLite\n"
            "作者：标书助手开发团队")


# ══════════════════════════════════════════════════════════════════
#  对话框：新建/编辑条目
# ══════════════════════════════════════════════════════════════════

